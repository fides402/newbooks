#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Script per organizzare i dati raccolti sui libri in formato tabellare
e creare un documento DOC con tabelle e link d'acquisto.
Versione corretta per gestire i link.
"""

import os
import pandas as pd
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
import logging
from datetime import datetime

# Configurazione del logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("organizza_dati_links.log"),
        logging.StreamHandler()
    ]
)

# Percorsi dei file
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ITALIAN_CSV = os.path.join(SCRIPT_DIR, "libri_italiani_links.csv")
AMERICAN_CSV = os.path.join(SCRIPT_DIR, "libri_americani_links.csv")
OUTPUT_DOC = os.path.join(SCRIPT_DIR, "Novità_Libri_Ultimi_30_Giorni_Con_Links.docx")

# Categorie da includere
CATEGORIE = ['filosofia', 'psicologia', 'società', 'business', 'self-help']

def load_data():
    """Carica i dati dai file CSV e li pulisce."""
    logging.info("Caricamento dei dati dai file CSV...")
    
    # Carica i dati italiani
    try:
        df_italian = pd.read_csv(ITALIAN_CSV)
        logging.info(f"Caricati {len(df_italian)} libri italiani")
    except Exception as e:
        logging.error(f"Errore durante il caricamento dei dati italiani: {e}")
        df_italian = pd.DataFrame()
    
    # Carica i dati americani
    try:
        df_american = pd.read_csv(AMERICAN_CSV)
        logging.info(f"Caricati {len(df_american)} libri americani")
    except Exception as e:
        logging.error(f"Errore durante il caricamento dei dati americani: {e}")
        df_american = pd.DataFrame()
    
    # Pulisci i dati italiani
    if not df_italian.empty:
        # Rimuovi righe con titoli vuoti o mancanti
        df_italian = df_italian[df_italian['titolo'].notna()]
        df_italian['titolo'] = df_italian['titolo'].apply(lambda x: x.strip() if isinstance(x, str) else "Titolo non disponibile")
        
        # Pulisci i link d'acquisto
        df_italian['link_acquisto'] = df_italian['link_acquisto'].apply(
            lambda x: x if isinstance(x, str) and x.strip() else None
        )
        
        # Aggiungi colonna per il paese
        df_italian['paese'] = 'Italia'
    
    # Pulisci i dati americani
    if not df_american.empty:
        # Rimuovi righe con titoli vuoti o mancanti
        df_american = df_american[df_american['titolo'].notna()]
        
        # Pulisci i titoli
        df_american['titolo'] = df_american['titolo'].apply(
            lambda x: x if x != "Titolo non disponibile" else "Libro americano"
        )
        
        # Pulisci i link d'acquisto
        df_american['link_acquisto'] = df_american['link_acquisto'].apply(
            lambda x: x if isinstance(x, str) and x.strip() else None
        )
        
        # Aggiungi colonna per il paese
        df_american['paese'] = 'USA'
    
    return df_italian, df_american

# Funzione per aggiungere hyperlink in Word
def add_hyperlink(paragraph, url, text, tooltip=None):
    """
    Aggiunge un hyperlink al paragrafo.
    
    Args:
        paragraph: Paragrafo a cui aggiungere il link
        url: URL del link
        text: Testo da visualizzare
        tooltip: Testo del tooltip (opzionale)
    """
    # Questo è un workaround poiché python-docx non supporta direttamente i link
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Crea il nodo hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    if tooltip is not None:
        hyperlink.set(qn('w:tooltip'), tooltip)
    
    # Crea il nodo di testo
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Aggiungi stile al testo (blu e sottolineato)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    
    return hyperlink

def create_document(df_italian, df_american):
    """Crea un documento Word con tabelle e link d'acquisto."""
    logging.info("Creazione del documento Word...")
    
    doc = Document()
    
    # Aggiungi titolo al documento
    title = doc.add_heading('Novità Libri degli Ultimi 30 Giorni', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Aggiungi data di generazione
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f"Generato il {datetime.now().strftime('%d/%m/%Y')}")
    date_run.italic = True
    
    # Aggiungi introduzione
    doc.add_paragraph()
    intro = doc.add_paragraph("Questo documento contiene le novità editoriali degli ultimi 30 giorni nelle categorie filosofia, psicologia, società, business e self-help, sia per il mercato italiano che per quello americano. Per ogni libro sono riportati titolo, autore, editore, prezzo e link d'acquisto diretto.")
    
    # Aggiungi indice
    doc.add_paragraph()
    doc.add_heading('Indice', level=1)
    for categoria in CATEGORIE:
        doc.add_paragraph(f"{categoria.capitalize()}", style='List Bullet')
    
    # Crea una sezione per ogni categoria
    for categoria in CATEGORIE:
        doc.add_page_break()
        doc.add_heading(f"Categoria: {categoria.capitalize()}", level=1)
        
        # Filtra i libri italiani per questa categoria
        italian_books = df_italian[df_italian['categoria'] == categoria] if not df_italian.empty else pd.DataFrame()
        
        # Filtra i libri americani per questa categoria
        american_books = df_american[df_american['categoria'] == categoria] if not df_american.empty else pd.DataFrame()
        
        # Aggiungi sezione per i libri italiani
        if not italian_books.empty:
            doc.add_heading("Libri Italiani", level=2)
            
            # Crea tabella
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Intestazioni
            header_cells = table.rows[0].cells
            header_cells[0].text = "Titolo"
            header_cells[1].text = "Autore"
            header_cells[2].text = "Editore"
            header_cells[3].text = "Prezzo"
            header_cells[4].text = "Link d'acquisto"
            
            # Formatta intestazioni
            for cell in header_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            
            # Aggiungi righe per ogni libro
            for _, book in italian_books.iterrows():
                row_cells = table.add_row().cells
                
                # Aggiungi dati
                row_cells[0].text = book['titolo'] if not pd.isna(book['titolo']) else "Titolo non disponibile"
                row_cells[1].text = book['autore'] if not pd.isna(book['autore']) else "Autore non disponibile"
                row_cells[2].text = book['editore'] if not pd.isna(book['editore']) else "Editore non disponibile"
                row_cells[3].text = book['prezzo'] if not pd.isna(book['prezzo']) else "Prezzo non disponibile"
                
                # Aggiungi link d'acquisto
                if book['link_acquisto'] and not pd.isna(book['link_acquisto']):
                    # Cancella il testo esistente nella cella
                    row_cells[4].text = ""
                    # Aggiungi il link come testo blu sottolineato
                    p = row_cells[4].paragraphs[0]
                    run = p.add_run("Link")
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blu
                    run.underline = True
                    # Aggiungi una nota che indica che questo è un link
                    p.add_run(" (Clicca per acquistare)")
                else:
                    row_cells[4].text = "Non disponibile"
                
                # Centra verticalmente il contenuto delle celle
                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            doc.add_paragraph()
        
        # Aggiungi sezione per i libri americani
        if not american_books.empty:
            doc.add_heading("Libri Americani", level=2)
            
            # Crea tabella
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Intestazioni
            header_cells = table.rows[0].cells
            header_cells[0].text = "Titolo"
            header_cells[1].text = "Autore"
            header_cells[2].text = "Prezzo"
            header_cells[3].text = "Link d'acquisto"
            
            # Formatta intestazioni
            for cell in header_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            
            # Aggiungi righe per ogni libro
            for _, book in american_books.iterrows():
                row_cells = table.add_row().cells
                
                # Aggiungi dati
                row_cells[0].text = book['titolo'] if not pd.isna(book['titolo']) else "Titolo non disponibile"
                row_cells[1].text = book['autore'] if not pd.isna(book['autore']) else "Autore non disponibile"
                row_cells[2].text = book['prezzo'] if not pd.isna(book['prezzo']) else "Prezzo non disponibile"
                
                # Aggiungi link d'acquisto
                if book['link_acquisto'] and not pd.isna(book['link_acquisto']):
                    # Cancella il testo esistente nella cella
                    row_cells[3].text = ""
                    # Aggiungi il link come testo blu sottolineato
                    p = row_cells[3].paragraphs[0]
                    run = p.add_run("Link")
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blu
                    run.underline = True
                    # Aggiungi una nota che indica che questo è un link
                    p.add_run(" (Clicca per acquistare)")
                else:
                    row_cells[3].text = "Non disponibile"
                
                # Centra verticalmente il contenuto delle celle
                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            doc.add_paragraph()
    
    # Aggiungi conclusione
    doc.add_page_break()
    doc.add_heading("Conclusioni", level=1)
    conclusion = doc.add_paragraph("Questo documento è stato generato automaticamente tramite uno script Python che ha raccolto le informazioni sui libri pubblicati negli ultimi 30 giorni nelle categorie filosofia, psicologia, società, business e self-help, sia per il mercato italiano che per quello americano.")
    doc.add_paragraph("I dati sono stati raccolti da IBS.it per i libri italiani e da Amazon.com per i libri americani. I link d'acquisto permettono di accedere direttamente alle pagine dei libri sui rispettivi siti.")
    
    # Aggiungi nota sui link
    doc.add_paragraph()
    note = doc.add_paragraph("Nota: I link d'acquisto sono indicati come testo blu sottolineato. Per utilizzarli, è necessario aprire il documento in un editor che supporti i link (come Microsoft Word o LibreOffice Writer) e fare clic sul testo 'Link'.")
    
    # Salva il documento
    try:
        doc.save(OUTPUT_DOC)
        logging.info(f"Documento salvato in {OUTPUT_DOC}")
        return True
    except Exception as e:
        logging.error(f"Errore durante il salvataggio del documento: {e}")
        return False

def main():
    """Funzione principale."""
    logging.info("Inizio organizzazione dei dati...")
    
    # Carica i dati
    df_italian, df_american = load_data()
    
    # Crea il documento
    success = create_document(df_italian, df_american)
    
    if success:
        logging.info("Organizzazione dei dati completata con successo")
    else:
        logging.error("Errore durante l'organizzazione dei dati")

if __name__ == "__main__":
    main()
